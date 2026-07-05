import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _add_bold_paragraph(doc, text: str, style_name: str = None):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    para = doc.add_paragraph(style=style_name)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    return para


def _set_paragraph_spacing(paragraph, before=0, after=6):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _process_paragraph(doc, line: str):
    stripped = line.strip()
    if not stripped:
        p = doc.add_paragraph("")
        _set_paragraph_spacing(p, before=0, after=2)
        return

    if stripped.startswith("---"):
        p = doc.add_paragraph("")
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        _set_paragraph_spacing(p, before=6, after=6)
        return

    if stripped.startswith("## "):
        heading = doc.add_heading(stripped[3:], level=2)
        _set_paragraph_spacing(heading, before=12, after=4)
        return

    if stripped.startswith("### "):
        heading = doc.add_heading(stripped[4:], level=3)
        _set_paragraph_spacing(heading, before=10, after=3)
        return

    if stripped.startswith("- ") or stripped.startswith("* "):
        text = stripped[2:]
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        _set_paragraph_spacing(p, before=0, after=2)
        return

    if re.match(r"^\d+[.)]\s", stripped):
        p = doc.add_paragraph(style="List Number")
        p.clear()
        text = re.sub(r"^\d+[.)]\s", "", stripped)
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        _set_paragraph_spacing(p, before=0, after=2)
        return

    _add_bold_paragraph(doc, stripped, style_name=None)
    para = doc.paragraphs[-1]
    _set_paragraph_spacing(para, before=0, after=6)


def _split_into_paragraphs(text: str):
    lines = text.split("\n")
    result = []
    current_para = []
    for line in lines:
        if line.strip() == "":
            if current_para:
                result.extend(current_para)
                current_para = []
            result.append("")
        else:
            current_para.append(line.rstrip())
    if current_para:
        result.extend(current_para)
    return result


def build_docx(title: str, sections: dict, output_dir: str) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(title_heading, before=0, after=18)

    for section_name, content in sections.items():
        section_heading = doc.add_heading(section_name, level=1)
        _set_paragraph_spacing(section_heading, before=16, after=8)

        paragraphs = _split_into_paragraphs(content)
        for line in paragraphs:
            if line.strip().lower() == section_name.lower():
                continue
            _process_paragraph(doc, line)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)
    safe_title = safe_title.replace(" ", "_").lower()[:50]
    filename = f"{safe_title}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)
    return filepath
